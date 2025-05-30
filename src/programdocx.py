#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sat Feb 11 08:49:33 2023

@author: repa
"""

# requires python-docx
from docx import Document

class WriteDocx:

    def __init__(self, project):

        self.pr = project

    def authorList(self, fname):

        doc = Document()
        authors = self.pr.author_list
        t = doc.add_table(len(authors), 2, style="Table Grid")
        for i, au in enumerate(authors):
            t.cell(i, 0).text = au.nameLastFirst()
            t.cell(i, 1).text = ', '.join(au.getEventCodes())

        doc.save(fname)

    def eventList(self, fname):

        doc = Document()
        for day in self.pr.getDays():

            doc.add_heading(day.printDate(), 2)

            for event in day.events:
                t = doc.add_table(2, 3, style="Table Grid")
                t.cell(0, 0).text = \
                    f'{event.printDay()} {event.printStart()} - {event.printEnd()}'
                t.cell(1, 0).text = event.venue
                t.cell(1, 1).text = event.printEventCode()
                t.cell(0, 1).text = event.printTitle()
                t.cell(1, 2).text = event.printChair()
                t.cell(0, 1).merge(t.cell(0,2))

                if event.hasSession():
                    for item in event._session._items:
                        r = t.add_row()
                        r.cells[0].merge(r.cells[2])
                        r.cells[0].add_paragraph(item.title)
                        r.cells[0].add_paragraph(item.printAuthors())
                doc.add_paragraph('')
        doc.save(fname)
